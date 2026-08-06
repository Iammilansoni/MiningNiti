"""
Document Text Extractors
Page-aware text extraction for PDF, DOCX, and plain-text files.

Replaces the monolithic _download_and_extract() in DocumentService with
dedicated extractor classes that track page boundaries — required for
context-aware answers that cite page numbers.
"""

import asyncio
import io
import logging
import os
import tempfile
from dataclasses import dataclass, field
from typing import List, Optional

import httpx

from app.config import settings
from app.core.url_guard import (
    fetch_remote_file,
    is_internal_storage_url,
    resolve_storage_path,
)

logger = logging.getLogger(__name__)


def _read_capped(path: str, max_bytes: int) -> bytes:
    """Read a local file, refusing anything over the size ceiling."""
    if os.path.getsize(path) > max_bytes:
        raise ValueError(
            f"Stored file exceeds the {max_bytes // (1024 * 1024)}MB limit"
        )
    with open(path, "rb") as f:
        return f.read()


# ── Data structures ────────────────────────────────────────────────────────────


@dataclass
class PageContent:
    """Text content of a single page."""

    page_number: int  # 1-indexed
    text: str  # raw text of this page
    char_start: int = 0  # character offset where this page starts in full_text
    char_end: int = 0  # character offset where this page ends


@dataclass
class ExtractedDocument:
    """Result of text extraction from a document file."""

    full_text: str
    pages: List[PageContent]
    total_pages: int
    file_type: str
    metadata: dict = field(default_factory=dict)  # author, title, creation_date, etc.

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


# ── Extractors ─────────────────────────────────────────────────────────────────


class PDFExtractor:
    """
    Page-aware PDF extraction.

    Primary path is pdfplumber, which additionally recovers tables (rendered
    as Markdown) and lets pages that yield no text be sent to OCR — both
    essential for mining documents, which are heavily tabular and frequently
    scanned. See app/services/pdf_layout.py.

    Falls back to plain pypdf if pdfplumber is unavailable or errors, so the
    ingest path never hard-fails on a library problem.
    """

    def extract(self, file_path: str) -> ExtractedDocument:
        try:
            return self._extract_with_layout(file_path)
        except ImportError:
            logger.info("pdfplumber not installed — using pypdf text-only path")
        except Exception as e:
            logger.warning(
                f"Layout-aware extraction failed ({e}); falling back to pypdf",
                exc_info=True,
            )
        return self._extract_with_pypdf(file_path)

    def _extract_with_layout(self, file_path: str) -> ExtractedDocument:
        """pdfplumber path: prose + Markdown tables + OCR for scanned pages."""
        import pdfplumber

        from app.services.pdf_layout import extract_page

        pages: List[PageContent] = []
        char_offset = 0
        total_tables = 0
        ocr_pages = 0
        ocr_budget = [settings.OCR_MAX_PAGES]

        with pdfplumber.open(file_path) as pdf:
            for page_num, plumber_page in enumerate(pdf.pages, start=1):
                result = extract_page(
                    plumber_page=plumber_page,
                    pdf_path=file_path,
                    page_number=page_num,
                    ocr_budget=ocr_budget,
                )
                total_tables += result.tables_found
                ocr_pages += 1 if result.used_ocr else 0

                pages.append(
                    PageContent(
                        page_number=page_num,
                        text=result.text,
                        char_start=char_offset,
                        char_end=char_offset + len(result.text),
                    )
                )
                char_offset += len(result.text)

            metadata = {}
            if pdf.metadata:
                for key, val in pdf.metadata.items():
                    if val:
                        metadata[str(key).lstrip("/").lower()] = str(val)

        full_text = "".join(p.text for p in pages)

        if total_tables or ocr_pages:
            logger.info(
                f"Layout extraction: {len(pages)} pages, "
                f"{total_tables} table(s), {ocr_pages} page(s) via OCR"
            )

        metadata["tables_extracted"] = total_tables
        metadata["ocr_pages"] = ocr_pages

        return ExtractedDocument(
            full_text=full_text,
            pages=pages,
            total_pages=len(pages),
            file_type="application/pdf",
            metadata=metadata,
        )

    def _extract_with_pypdf(self, file_path: str) -> ExtractedDocument:
        from pypdf import PdfReader

        pages: List[PageContent] = []
        char_offset = 0

        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                pages.append(
                    PageContent(
                        page_number=page_num,
                        text=page_text,
                        char_start=char_offset,
                        char_end=char_offset + len(page_text),
                    )
                )
                char_offset += len(page_text)

            full_text = "".join(p.text for p in pages)

            # Extract PDF metadata
            metadata = {}
            if reader.metadata:
                for key, val in reader.metadata.items():
                    if val:
                        clean_key = key.lstrip("/").lower()
                        metadata[clean_key] = str(val)

            return ExtractedDocument(
                full_text=full_text,
                pages=pages,
                total_pages=len(pages),
                file_type="application/pdf",
                metadata=metadata,
            )

        except Exception as e:
            logger.error(f"pypdf extraction failed: {e}", exc_info=True)
            raise ValueError(f"Failed to extract text from PDF: {e}")


class DocxExtractor:
    """DOCX text extraction with paragraph tracking (approximate page numbers)."""

    def extract(self, file_path: str) -> ExtractedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n".join(paragraphs)

        # DOCX doesn't expose true page numbers in python-docx.
        # We approximate: every ~3000 chars ≈ 1 page.
        pages = self._approximate_pages(full_text)

        return ExtractedDocument(
            full_text=full_text,
            pages=pages,
            total_pages=len(pages),
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={},
        )

    def _approximate_pages(
        self, text: str, chars_per_page: int = 3000
    ) -> List[PageContent]:
        pages = []
        for i in range(0, max(1, len(text)), chars_per_page):
            page_text = text[i : i + chars_per_page]
            pages.append(
                PageContent(
                    page_number=len(pages) + 1,
                    text=page_text,
                    char_start=i,
                    char_end=i + len(page_text),
                )
            )
        return pages or [
            PageContent(page_number=1, text=text, char_start=0, char_end=len(text))
        ]


class PlainTextExtractor:
    """Plain text extraction."""

    def extract(self, file_path: str) -> ExtractedDocument:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()

        # Approximate pages for plain text too
        chars_per_page = 3000
        pages = []
        for i in range(0, max(1, len(full_text)), chars_per_page):
            page_text = full_text[i : i + chars_per_page]
            pages.append(
                PageContent(
                    page_number=len(pages) + 1,
                    text=page_text,
                    char_start=i,
                    char_end=i + len(page_text),
                )
            )

        return ExtractedDocument(
            full_text=full_text,
            pages=pages
            or [
                PageContent(
                    page_number=1, text=full_text, char_start=0, char_end=len(full_text)
                )
            ],
            total_pages=max(1, len(pages)),
            file_type="text/plain",
        )


# ── File downloader + dispatcher ───────────────────────────────────────────────

EXTENSION_MAP = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "text/plain": ".txt",
}

_extractors = {
    "pdf": PDFExtractor(),
    "docx": DocxExtractor(),
    "txt": PlainTextExtractor(),
}


async def download_and_extract(file_url: str, file_type: str) -> ExtractedDocument:
    """
    Read a document's bytes and extract its text content with page tracking.

    Two sources are supported, and they are deliberately not interchangeable:

      storage:// (or legacy file://)  — a file this backend stored itself.
                                        Resolved inside UPLOAD_DIR only.
      https://                        — a remote URL. Passed through the SSRF
                                        guard, size-capped, redirects revalidated.

    Args:
        file_url: Internal storage URL or a public https URL
        file_type: MIME type string

    Returns:
        ExtractedDocument with full text and per-page content

    Raises:
        UnsafeURLError: if the URL is not fetchable safely
    """
    suffix = EXTENSION_MAP.get(file_type, ".tmp")
    max_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024

    if is_internal_storage_url(file_url):
        # Path is confined to UPLOAD_DIR by the guard; traversal is rejected.
        local_path = await asyncio.to_thread(resolve_storage_path, file_url)
        file_bytes = await asyncio.to_thread(_read_capped, local_path, max_bytes)
    else:
        # User-supplied URL: https only, public addresses only, size-capped.
        file_bytes = await fetch_remote_file(file_url, max_bytes=max_bytes)

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        if "pdf" in file_type.lower():
            return _extractors["pdf"].extract(tmp_path)
        elif "docx" in file_type.lower() or "wordprocessingml" in file_type.lower():
            return _extractors["docx"].extract(tmp_path)
        else:
            return _extractors["txt"].extract(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
